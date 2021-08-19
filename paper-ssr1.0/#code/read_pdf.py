import pdfplumber
from openpyxl import load_workbook
from defination import Load_number_ExcelDone
import docx
from defination import judge_filename
def save_word(text_list,n,filepath):
    try:doc=docx.Document(filepath)
    except:doc=docx.Document()
    doc.add_paragraph('---------------------------------'+str(n)+'-----------------------------------------')
    for text in text_list:
        doc.add_paragraph(text)
    doc.save(filepath)
    
def readtext_pdf(path):
    with pdfplumber.open(path) as pdf:
        page=pdf.pages[0]
        Text=page.extract_text()
    return Text

def judge_number(start,end):
    class index:
        def __init__(self,start,end):
            self.start=start
            self.end=end
    judge_e=0
    judge_s=0
    while papertext[end-1].isdigit() or papertext[end+1].isdigit():
        end=papertext.find('.',end+1)
        judge_e=1+judge_e
        # print(1)
        if judge_e>20:
            break
    while papertext[start-1].isdigit() or papertext[start+1].isdigit():
        start=papertext.rfind('.',0,start-1)
        judge_s=1+judge_s
        # print(2)
        if judge_s>20:
           break
    start=start+1
    result=index(start,end)
    return result

def add_usefulltext_excel(filepath,n,list):
    workbook=load_workbook(filepath+'.xlsx')
    wb=workbook.active
    count=70
    for sentence in list:
        count=count+1
        column=chr(count)
        row=str(n)
        cell_location=column+row
        wb[cell_location]=sentence
    workbook.save(filepath+'.xlsx')

def get_data_from_sentense(sentence_list):
    lists=['.','-','—','–',' ','−']
    for sentence in sentence_list:
        sentence=sentence.replace('c','').replace('i','').replace('d','').replace('(','').replace(')','').replace(':','').replace('i','')
        list_data=[]
        list_number=[]
        ele='None'
        for i,s in enumerate(sentence):
            if s.isdigit() or s in lists:
                list_number.append(s)
            if s=='N' and sentence[i+1]=='d':  
                ele='Nd:'
            elif s=='H' and sentence[i+1]=='f':
                ele='Hf:'

        num=''.join(list_number)
        data=ele+num
        list_data.append(data)
    return list_data
                  
number_list=[]
data_list=[]
Load_number_ExcelDone('./excel/paper.xlsx',number_list)
for n in number_list:
    try:    
        papertext_usefull=[]
        field_n=judge_filename(n)
        # print(field_n)
        path='./paper_pdf/'+field_n+str(n)+'.pdf'
        papertext=readtext_pdf(path)
        for i,s in enumerate(papertext):
            if s=='ε':
                end_index=papertext.find('.',i)
                start_index=papertext.rfind('.',0,i)
                result_numberjudge=judge_number(start_index,end_index)
                end_index=result_numberjudge.end
                start_index=result_numberjudge.start
                usefulltext=papertext[start_index:end_index]
                usefulltext=usefulltext.replace('\r','').replace('\n','').replace('\t','')
                print(usefulltext)
                papertext_usefull.append(usefulltext)

        data=get_data_from_sentense(papertext_usefull)
        for da in data:
            print('==============================='+da+'===========================================')
        save_word(papertext_usefull,n,'./text.docx')
        add_usefulltext_excel('./excel/paper',n,data)
    except:
        print('________________________'+str(n)+'load failed________________________________')





